import os
import time
import re
from dataclasses import dataclass
from datetime import datetime

from docx import Document
from docx.document import Document as _Document
from docx.table import _Cell, Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph

from openai import OpenAI

# ================== 基本配置 ==================

# ===================== 基本配置（开源安全版） =====================

DEFAULT_BASE_DIR = ""  # 开源版默认不写死路径；需要时可填绝对路径

if DEFAULT_BASE_DIR and os.path.isdir(DEFAULT_BASE_DIR):
    BASE_DIR = DEFAULT_BASE_DIR
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))  # __file__ 为当前脚本路径 :contentReference[oaicite:0]{index=0}

ENABLE_LOGO = False  # 开源版默认不启用 Logo
LOGO_PATH = os.path.join(BASE_DIR, "assets", "logo.png")  # 中性路径

SOURCE_DIR = os.path.join(BASE_DIR, "source_cn")
OUTPUT_DIR = os.path.join(BASE_DIR, "translated_en")

TRANSLATOR_NAME = ""  # 开源版默认空；留空即不写署名


MODEL_NAME = "deepseek-chat"
BASE_URL = "https://api.deepseek.com"

REQUEST_INTERVAL = 1.0       # 每次请求之间的最小间隔（秒）
MAX_RETRY_PER_BLOCK = 3      # 单个段落/单元格最大重试次数
MAX_CONSEC_FAILS = 10        # 整个文件允许的最大连续“硬失败”次数


client = None

# ================== 术语表 ==================

TERMINOLOGY_MAP = {
    "不合格品控制程序": "Control of Nonconforming Product Procedure",
    "不合格品处理单": "Nonconforming Product Disposition Form",
    "不合格品": "Nonconforming Product",
    "不合格": "Nonconformity",
    "让步接收": "Concession Acceptance",
    "返工": "Rework",
    "返修": "Repair",
    "冻结": "Hold",
    "报废": "Scrap",
    "物料评审委员会": "Material Review Board (MRB)",
    "MRB评审": "MRB Review",
    "进料检验": "Incoming Inspection",
    "过程检验": "In-process Inspection",
    "最终检验": "Final Inspection",
    "抽检": "Sampling Inspection",
    "检验记录": "Inspection Record",
    "处置": "Disposition",
    "处置意见": "Disposition Decision",
    "纠正措施": "Corrective Action",
    "预防措施": "Preventive Action",
    "纠正和预防措施": "Corrective and Preventive Action",

    "non-conforming product": "Nonconforming Product",
    "Non-conforming product": "Nonconforming Product",
    "nonconforming product": "Nonconforming Product",
    "nonconforming products": "Nonconforming Products",
    "Nonconforming products": "Nonconforming Products",

    "material review board": "Material Review Board (MRB)",
    "Material Review Board": "Material Review Board (MRB)",

    "incoming inspection": "Incoming Inspection",
    "in-process inspection": "In-process Inspection",
    "final inspection": "Final Inspection",
    "sampling inspection": "Sampling Inspection",

    "quality management system": "Quality Management System",
    "qms": "QMS",
    "supplier": "Supplier",
    "suppliers": "Suppliers",
}


# ================== 外部术语表（glossary_dict.py，可选） ==================

# 说明：如果 glossary_dict.py 存在，则自动加载其中的 GLOSSARY_CN_TO_EN，
# 后续术语表会更新
# 并在翻译提示词中强制模型采用这些“首选译法”，以提高一致性与专业度。
# 注意：glossary_dict.py 是翻译工具文件，不属于受控 QMS 文档；用于锁定术语输出。

GLOSSARY_CN_TO_EN = {}
GLOSSARY_PROMPT_BLOCK = ""

def load_external_glossary():
    global GLOSSARY_CN_TO_EN, GLOSSARY_PROMPT_BLOCK

    # 查找 glossary_dict.py 文件
    candidates = [
        os.path.join(BASE_DIR, "glossary_dict.py"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "glossary_dict.py"),
    ]

    for p in candidates:
        if os.path.isfile(p):
            try:
                import importlib.util
                spec = importlib.util.spec_from_file_location("glossary_dict", p)
                mod = importlib.util.module_from_spec(spec)
                assert spec and spec.loader
                spec.loader.exec_module(mod)

                if hasattr(mod, "GLOSSARY_CN_TO_EN") and isinstance(mod.GLOSSARY_CN_TO_EN, dict):
                    GLOSSARY_CN_TO_EN = mod.GLOSSARY_CN_TO_EN
                    # 生成可直接塞进 system prompt 的术语块（控制长度，避免提示词过大）
                    lines = []
                    # 最长优先，减少部分包含关系导致的歧义
                    for cn in sorted(GLOSSARY_CN_TO_EN.keys(), key=len, reverse=True):
                        en = GLOSSARY_CN_TO_EN[cn]
                        # 只收录短词条；过长句子容易让模型“展开解释” 节约token哈哈哈哈哈
                        if len(cn) <= 40 and len(en) <= 80:
                            lines.append(f"- {cn} => {en}")
                    if lines:
                        GLOSSARY_PROMPT_BLOCK = "Preferred terminology (use exactly as specified):\n" + "\n".join(lines) + "\n"
                    print(f"[INFO] 已加载外部术语表: {p}（条目数: {len(GLOSSARY_CN_TO_EN)}）")
                    return
            except Exception as e:
                print(f"[WARN] 外部术语表加载失败: {p}，将继续使用内置术语表。原因：{e}")
                return

    # 未找到则保持为空
    print("[INFO] 未找到 glossary_dict.py，继续使用内置术语表。")


# ================== 工具函数 ==================

def init_client():
    global client
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise RuntimeError(
            "未找到环境变量 DEEPSEEK_API_KEY，请在 PowerShell 中先执行：\n"
            '$env:DEEPSEEK_API_KEY="你的 DeepSeek API Key"'
        )
    client = OpenAI(
        api_key=api_key,
        base_url=BASE_URL,
    )

def has_chinese(text: str) -> bool:
    if not text:
        return False
    for ch in text:
        if "\u4e00" <= ch <= "\u9fff":
            return True
    return False

def unify_terms(text: str) -> str:
    if not text:
        return text
    result = text
    # 最长优先，避免短词条先替换导致的错误覆盖
    for k in sorted(TERMINOLOGY_MAP.keys(), key=len, reverse=True):
        result = result.replace(k, TERMINOLOGY_MAP[k])
    return result

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("iter_block_items: 不支持的 parent 类型")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def count_translatable_blocks(doc: Document) -> int:
    total = 0
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            txt = block.text.strip()
            if txt and has_chinese(txt):
                total += 1
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt and has_chinese(txt):
                        total += 1
    return total

def format_minutes(seconds: float) -> str:
    if seconds <= 0:
        return "少于 1 分钟"
    minutes = seconds / 60.0
    if minutes < 1:
        return "少于 1 分钟"
    return f"约 {minutes:.1f} 分钟"

def print_progress(filename: str, done_blocks: int, total_blocks: int, start_time: float):
    if total_blocks <= 0:
        return
    ratio = done_blocks / total_blocks
    ratio = max(0.0, min(1.0, ratio))
    percent = ratio * 100

    bar_len = 30
    filled = int(bar_len * ratio)
    bar = "#" * filled + "-" * (bar_len - filled)

    elapsed = time.time() - start_time
    if done_blocks > 0:
        est_total = elapsed / done_blocks * total_blocks
        remain = max(0.0, est_total - elapsed)
    else:
        remain = 0.0

    msg = (
        f"[PROGRESS] {filename} [{bar}] "
        f"{percent:5.1f}% 块 {done_blocks}/{total_blocks} "
        f"预计剩余时间：{format_minutes(remain)}"
    )
    print("\r" + msg, end="", flush=True)

# ================== 译文风控：只做软告警，不回退中文 ==================

@dataclass
class SegmentQuality:
    ok: bool
    warning: str | None = None
    suspicious: bool = False

def _is_probably_title(cn: str) -> bool:
    s = (cn or "").strip()
    if not s:
        return True
    if len(s) <= 12:
        return True
    if re.fullmatch(r"[0-9一二三四五六七八九十]+[\.、．]\s*.*", s):
        return True
    if "：" in s and len(s) <= 30:
        return True
    return False

def _safe_len_ratio_threshold(cn: str) -> float:
    n = len((cn or "").strip())
    if n <= 6:
        return 6.0
    if n <= 20:
        return 4.5
    if n <= 60:
        return 4.0
    return 3.5

def assess_translation_quality(cn: str, en: str) -> SegmentQuality:
    cn_s = (cn or "").strip()
    en_s = (en or "").strip()

    if not cn_s:
        return SegmentQuality(ok=True)

    if not en_s:
        return SegmentQuality(ok=False, warning="[WARN] 译文为空")

    ratio = len(en_s) / max(1, len(cn_s))
    thr = _safe_len_ratio_threshold(cn_s)

    if _is_probably_title(cn_s):
        thr = max(thr, 8.0)

    if ratio > thr:
        return SegmentQuality(
            ok=True,
            warning=(
                f"[WARN] 译文长度偏长（原文 {len(cn_s)} 字，译文 {len(en_s)} 字，ratio={ratio:.2f} > 阈值 {thr:.2f}）。"
                f"不终止流程，仅标记需人工复核。"
            ),
            suspicious=True
        )

    if ratio < 0.35 and len(cn_s) >= 12:
        return SegmentQuality(
            ok=True,
            warning=(
                f"[WARN] 译文长度偏短（原文 {len(cn_s)} 字，译文 {len(en_s)} 字，ratio={ratio:.2f}）。"
                f"不终止流程，仅标记需人工复核。"
            ),
            suspicious=True
        )

    return SegmentQuality(ok=True)

def normalize_fail_category(err: Exception | str) -> str:
    msg = str(err).lower()
    if any(k in msg for k in ["timeout", "timed out", "read timeout", "connect timeout"]):
        return "network"
    if any(k in msg for k in ["429", "rate limit", "too many requests"]):
        return "rate_limit"
    if any(k in msg for k in ["500", "502", "503", "504", "bad gateway", "service unavailable"]):
        return "server"
    if any(k in msg for k in ["json", "parse", "decode"]):
        return "parse"
    return "other"

# ================== 调用 DeepSeek 翻译 ==================

def translate_text_once(cn: str) -> str:
    if not cn.strip():
        return cn

    global client
    if client is None:
        init_client()

    system_prompt = (
        "You are a professional translator for ISO 9001 factory Quality Management System (QMS) documents. "
        "Your ONLY task is to translate existing Chinese text into English.\n"
        "Requirements:\n"
        "1) Translate faithfully, sentence by sentence.\n"
        "2) Preserve the original structure, numbering, and bullet points.\n"
        "3) DO NOT add new clauses, sections, examples, or explanations.\n"
        "4) DO NOT summarize or rewrite; only translate.\n"
        "5) Use concise, formal American English suitable for audited QMS procedures.\n"
        "6) Output English only, no Chinese.\n"
    )


    # 强制术语一致性（若已加载外部术语表）
    if GLOSSARY_PROMPT_BLOCK:
        system_prompt += "\n" + GLOSSARY_PROMPT_BLOCK
    user_prompt = (
        "Translate the following Chinese QMS text into English.\n"
        "The output must keep roughly the same structure and level of detail as the input.\n\n"
        f"Chinese text:\n{cn}"
    )

    resp = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.0,
    )
    content = resp.choices[0].message.content or ""
    return content.strip()

def translate_with_retries(cn: str) -> tuple[str, Exception | None]:
    last_error = None
    for attempt in range(1, MAX_RETRY_PER_BLOCK + 1):
        try:
            en = translate_text_once(cn)
            return en, None
        except Exception as e:
            last_error = e
            cat = normalize_fail_category(e)
            print(f"\n[WARN] 翻译请求失败（{cat}），第 {attempt}/{MAX_RETRY_PER_BLOCK} 次重试：{e}")
            time.sleep(2.0)
    return "", last_error

def process_translation_segment(cn: str) -> tuple[str, bool, bool]:
    """
    返回: (text_to_write, hard_failed, suspicious)
    hard_failed=True 只用于累计连续失败并触发 FATAL
    """
    en, err = translate_with_retries(cn)

    if err is not None:
        print(f"\n[ERROR] 段落翻译失败，连续失败 {MAX_RETRY_PER_BLOCK} 次：{err}")
        return cn, True, True

    if not en.strip():
        print(f"\n[WARN] 译文为空，本段保留中文原文，需人工处理。")
        return cn, True, True

    en = unify_terms(en)

    if has_chinese(en):
        print(f"\n[WARN] 译文仍包含中文字符，判定输出不合格，本段保留中文原文，需人工处理。")
        return cn, True, True

    quality = assess_translation_quality(cn, en)
    if quality.warning:
        print("\n" + quality.warning)

    return en, False, quality.suspicious

# ================== 文档翻译逻辑 ==================

def insert_translator_info(doc: Document):
    # 开源版默认不写署名；需要时可把 TRANSLATOR_NAME 设置为“某某事业部翻译”
    if not TRANSLATOR_NAME:
        return

    para0 = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    info = f"Translated by {TRANSLATOR_NAME}. Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}."
    para0.insert_paragraph_before(info)

def translate_document(in_path: str, out_path: str):
    filename = os.path.basename(in_path)
    print(f"\n[INFO] 开始翻译文件: {filename}")

    doc = Document(in_path)

    total_blocks = count_translatable_blocks(doc)
    if total_blocks == 0:
        print("[INFO] 文件中没有需要翻译的中文，直接复制保存。")
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        doc.save(out_path)
        print(f"[OK] 输出英文文件: {out_path}")
        return

    print(f"[INFO] 需要翻译的段落/单元格总数: {total_blocks}")
    start_time = time.time()
    done_blocks = 0
    consecutive_hard_failures = 0
    last_request_time = 0.0
    suspicious_segments = 0

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            txt = block.text.strip()
            if not txt or not has_chinese(txt):
                continue

            now = time.time()
            delta = now - last_request_time
            if delta < REQUEST_INTERVAL:
                time.sleep(REQUEST_INTERVAL - delta)
            last_request_time = time.time()

            cn = block.text
            out_text, hard_failed, suspicious = process_translation_segment(cn)

            if suspicious:
                suspicious_segments += 1

            if hard_failed:
                consecutive_hard_failures += 1
            else:
                consecutive_hard_failures = 0

            if consecutive_hard_failures >= MAX_CONSEC_FAILS:
                print(
                    f"\n[FATAL] 文件 {filename} 连续硬失败已达 {MAX_CONSEC_FAILS} 段，"
                    f"判定接口/输出持续异常，终止本文件翻译，不保存输出。"
                )
                return

            block.text = out_text

            done_blocks += 1
            print_progress(filename, done_blocks, total_blocks, start_time)

        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if not txt or not has_chinese(txt):
                        continue

                    now = time.time()
                    delta = now - last_request_time
                    if delta < REQUEST_INTERVAL:
                        time.sleep(REQUEST_INTERVAL - delta)
                    last_request_time = time.time()

                    cn = cell.text
                    out_text, hard_failed, suspicious = process_translation_segment(cn)

                    if suspicious:
                        suspicious_segments += 1

                    if hard_failed:
                        consecutive_hard_failures += 1
                    else:
                        consecutive_hard_failures = 0

                    if consecutive_hard_failures >= MAX_CONSEC_FAILS:
                        print(
                            f"\n[FATAL] 文件 {filename} 连续硬失败已达 {MAX_CONSEC_FAILS} 段，"
                            f"判定接口/输出持续异常，终止本文件翻译，不保存输出。"
                        )
                        return

                    cell.text = out_text

                    done_blocks += 1
                    print_progress(filename, done_blocks, total_blocks, start_time)

    print()
    if suspicious_segments > 0:
        print(f"[INFO] 本文件需人工复核段落数（软告警/保留中文等）：{suspicious_segments}")

    insert_translator_info(doc)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"[OK] 完成翻译并保存: {out_path}")

# ================== 主入口 ==================

def ensure_dirs():
    os.makedirs(SOURCE_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"[INFO] 源文件目录: {SOURCE_DIR}")
    print(f"[INFO] 输出目录: {OUTPUT_DIR}")
    if not os.path.exists(LOGO_PATH):
        print(f"[WARN] 未找到 logo 文件: {LOGO_PATH}，当前未启用 logo 插入。")

    # 尝试加载外部术语表（若存在）
    load_external_glossary()

def main():
    ensure_dirs()

    files = [f for f in os.listdir(SOURCE_DIR) if f.lower().endswith(".docx")]
    if not files:
        print("[INFO] 源目录中没有找到 .docx 文件，请将要翻译的中文体系文件放到 source_cn 目录。")
        return

    print(f"[INFO] 将翻译以下文件（目录: {SOURCE_DIR}）:")
    for name in files:
        print("  ", name)

    for name in files:
        in_path = os.path.join(SOURCE_DIR, name)
        out_name = os.path.splitext(name)[0] + "_EN.docx"
        out_path = os.path.join(OUTPUT_DIR, out_name)

        try:
            translate_document(in_path, out_path)
        except Exception as e:
            print(f"\n[ERROR] 文件 {name} 翻译过程中发生未处理异常，本文件跳过：{e}")

if __name__ == "__main__":
    main()
