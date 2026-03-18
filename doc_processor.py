import os
import logging
from typing import List
from docx import Document


class DocProcessor:
    def __init__(self, translator_name: str = "Kerrmote Yao"):
        self.translator_name = translator_name
        self.engine_name = "DeepSeek (深度求索)"
        self.logger = logging.getLogger("DocProcessor")

    def read_docx(self, path: str) -> List[str]:
        doc = Document(path)
        return [p.text for p in doc.paragraphs if p.text.strip()]

    def translate_docx(self, input_path: str, output_path: str, engine, direction: str, stop_event):
        doc = Document(input_path)

        # 翻译段落
        for p in doc.paragraphs:
            if stop_event and stop_event.is_set():
                return False
            if p.text.strip():
                translated = engine.translate_segment(p.text, direction, stop_event)
                if not translated and p.text.strip():
                    raise ValueError(f"段落翻译失败，未生成译文: {p.text[:50]}...")
                p.text = translated

        # 翻译表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if stop_event and stop_event.is_set():
                            return False
                        if p.text.strip():
                            translated = engine.translate_segment(p.text, direction, stop_event)
                            if not translated and p.text.strip():
                                raise ValueError("表格单元格翻译失败，未生成译文")
                            p.text = translated

        # 只有在没有停止且翻译成功的情况下才保存
        if not (stop_event and stop_event.is_set()):
            doc.save(output_path)
            return True
        return False

    def read_txt(self, path: str) -> List[str]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return [line.strip() for line in f.readlines() if line.strip()]

    def save_txt(self, path: str, content: List[str]):
        if not content:
            raise ValueError("尝试保存空的 TXT 文件")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(content))

    def translate_file(self, input_path: str, output_dir: str, engine) -> str:
        """
        Translate a single file and return the output file path.

        Supported: .docx, .txt
        Direction/stop_event are taken from engine when available to keep GUI calls simple.
        """
        if not os.path.isfile(input_path):
            raise FileNotFoundError(f"找不到输入文件: {input_path}")

        ext = os.path.splitext(input_path)[1].lower()

        # Normalize direction: accept zh2en/en2zh and zh-en/en-zh
        direction = getattr(engine, "direction", "zh-en")
        if isinstance(direction, str):
            direction = direction.replace("2", "-")
        else:
            direction = "zh-en"

        stop_event = getattr(engine, "stop_event", None)

        os.makedirs(output_dir, exist_ok=True)

        base = os.path.splitext(os.path.basename(input_path))[0]
        suffix = "_EN" if direction == "zh-en" else "_ZH"
        out_path = os.path.join(output_dir, f"{base}{suffix}{ext}")

        if ext == ".docx":
            ok = self.translate_docx(input_path, out_path, engine, direction, stop_event)
            if not ok:
                raise RuntimeError("任务被停止，未输出文件。")
            return out_path

        if ext == ".txt":
            lines = self.read_txt(input_path)
            out_lines = []
            for line in lines:
                if stop_event and stop_event.is_set():
                    raise RuntimeError("任务被停止，未输出文件。")
                out_lines.append(engine.translate_segment(line, direction, stop_event))
            self.save_txt(out_path, out_lines)
            return out_path

        raise ValueError(f"暂不支持该文件类型: {ext}（仅支持 .docx / .txt）")
